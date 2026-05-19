"""
Generate the weekly Word document from articles across multiple journals.

Each journal gets its own section with a 7-column table:
  # / Title (EN) / Title (ZH) / Authors / Abstract (EN) / Abstract (ZH) / Date
"""

import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger(__name__)

COLOR_PRIMARY = RGBColor(0x1F, 0x38, 0x64)
COLOR_ACCENT = RGBColor(0x2E, 0x75, 0xB6)
COLOR_GREY = RGBColor(0x59, 0x59, 0x59)


def _set_landscape(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)


def _add_title_block(doc, journal_results, days_back):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Multi-Journal Weekly Research Digest")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = COLOR_PRIMARY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("多期刊每周研究文章汇总")
    run.italic = True
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_GREY

    total = sum(len(v) for v in journal_results.values())
    today = datetime.now().strftime("%Y-%m-%d")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        f"Generated: {today}  |  Total: {total} articles  |  "
        f"Journals: {len(journal_results)}  |  "
        f"Window: past {days_back} days"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_ACCENT

    # Summary table of counts per journal
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Articles per journal")
    run.bold = True
    run.font.size = Pt(11)

    summary_table = doc.add_table(rows=1 + len(journal_results), cols=2)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_table.style = "Light Grid Accent 1"

    # Header
    hdr = summary_table.rows[0].cells
    for i, h in enumerate(["Journal", "# Articles"]):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2E75B6")
        hdr[i]._tc.get_or_add_tcPr().append(shd)

    for row_idx, (name, articles) in enumerate(journal_results.items(), start=1):
        row = summary_table.rows[row_idx].cells
        row[0].text = ""
        run = row[0].paragraphs[0].add_run(name)
        run.font.size = Pt(10)
        row[1].text = ""
        p = row[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(len(articles)))
        run.font.size = Pt(10)
        run.bold = True


def _add_section_heading(doc, text, page_break=True):
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = COLOR_PRIMARY


def _add_article_table(doc, articles):
    """Add a 7-column table with article metadata + Chinese translations."""
    if not articles:
        p = doc.add_paragraph()
        p.add_run(
            "No new research articles found in this time window."
        ).italic = True
        return

    table = doc.add_table(rows=1 + len(articles), cols=7)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    col_widths = [
        Cm(0.8), Cm(4.0), Cm(4.0), Cm(3.2),
        Cm(6.5), Cm(6.5), Cm(1.7),
    ]
    headers = ["#", "Title (EN)", "标题（中）", "Authors",
               "Abstract (EN)", "摘要（中）", "Date"]

    # Header row
    for col_idx, (header, width) in enumerate(zip(headers, col_widths)):
        cell = table.rows[0].cells[col_idx]
        cell.width = width
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2E75B6")
        cell._tc.get_or_add_tcPr().append(shd)

    for i, article in enumerate(articles, start=1):
        row = table.rows[i]
        cells = row.cells
        for c, w in zip(cells, col_widths):
            c.width = w

        cells[0].text = str(i)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        cells[1].text = ""
        run = cells[1].paragraphs[0].add_run(article.get("title", ""))
        run.bold = True
        run.font.size = Pt(9)

        cells[2].text = ""
        run = cells[2].paragraphs[0].add_run(article.get("title_zh", ""))
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_PRIMARY

        cells[3].text = ""
        run = cells[3].paragraphs[0].add_run(article.get("authors", ""))
        run.italic = True
        run.font.size = Pt(9)

        abstract_en = article.get("abstract") or ""
        if not abstract_en.strip():
            abstract_en = "[Abstract not yet available — see DOI link]"
        elif len(abstract_en) > 1500:
            abstract_en = abstract_en[:1500].rstrip() + "..."
        cells[4].text = ""
        run = cells[4].paragraphs[0].add_run(abstract_en)
        run.font.size = Pt(9)

        abstract_zh = article.get("abstract_zh") or ""
        if not abstract_zh.strip():
            abstract_zh = "[摘要暂未获取到，请点击 DOI 查看原文]"
        elif len(abstract_zh) > 1500:
            abstract_zh = abstract_zh[:1500].rstrip() + "..."
        cells[5].text = ""
        run = cells[5].paragraphs[0].add_run(abstract_zh)
        run.font.size = Pt(9)

        cells[6].text = ""
        p = cells[6].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(article.get("pub_date", ""))
        run.font.size = Pt(8)
        if article.get("doi"):
            p2 = cells[6].add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p2.add_run(article["doi"])
            run.font.size = Pt(6)
            run.font.color.rgb = COLOR_GREY


def add_classification_section(doc, report):
    """Append the per-category summary section at the end of the document."""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("📊 本期文献分类总结")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = COLOR_PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Classification Summary")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_GREY

    p = doc.add_paragraph()
    run = p.add_run(
        f"本期共收录 {report['total_articles']} 篇研究文章，"
        f"按学科自动分入 {len(report['categories'])} 个类别。"
        f"由于现代研究普遍跨学科，同一篇文章可能出现在多个类别中。"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_GREY
    doc.add_paragraph()

    for category in report["categories"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"▌ {category['name']}（{len(category['articles'])} 篇）")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_ACCENT

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(category["summary"])
        run.font.size = Pt(10)

        for art in category["articles"]:
            journal = art.get("_journal") or art.get("journal", "")
            title_en = art.get("title", "")
            title_zh = art.get("title_zh", "")

            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"[{journal}] ")
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_ACCENT
            run = p.add_run(title_en)
            run.font.size = Pt(9)

            if title_zh:
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Cm(1.0)
                p2.paragraph_format.space_after = Pt(4)
                run = p2.add_run(f"→ {title_zh}")
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = COLOR_GREY


def _add_footer(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Data sources: Crossref REST API + nature.com + Semantic Scholar API. "
        "Abstracts © respective publishers; reproduced for fair-use review only."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_GREY


def create_digest_document(
    journal_results,
    output_path,
    days_back=8,
    classification_report=None,
):
    """Build the multi-journal .docx file and return its path.

    Args:
        journal_results: dict {journal_name: [article, ...]} from fetch_all_journals()
        output_path: where to save the .docx file
        days_back: time window (just for display)
        classification_report: optional output from build_classification_report()
    """
    doc = Document()
    _set_landscape(doc)
    _add_title_block(doc, journal_results, days_back)

    # One section per journal, in the order they appear in journal_results
    for journal_name, articles in journal_results.items():
        _add_section_heading(doc, f"{journal_name} — Research Articles ({len(articles)})",
                             page_break=True)
        _add_article_table(doc, articles)

    if classification_report and classification_report.get("categories"):
        add_classification_section(doc, classification_report)

    _add_footer(doc)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info("Document saved: %s", output_path)
    return output_path
